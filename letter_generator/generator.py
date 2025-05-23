"""Main letter generation logic."""
import logging
import os
import time
import subprocess
from datetime import datetime
from pathlib import Path
from typing import Optional, Dict
from num2words import num2words
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENTATION
from docx2pdf import convert
import win32com.client
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

from constants import get_asset_path
from .exceptions import GenerationError
from .templates import (
    COMPANY_FOOTER, ANNUAL_LETTER_TEMPLATE,
    FIFTEEN_YEAR_LETTER_TEMPLATE, SECOND_LETTER_TEMPLATE
)
from .formatter import format_names, format_address, generate_filename
from .document_processor import (
    extract_names_and_address_annual,
    extract_names_and_address_fifteen_year,
    validate_content
)

logger = logging.getLogger(__name__)

# Try to import psutil, but don't fail if it's not available
try:
    import psutil
    PSUTIL_AVAILABLE = True
except ImportError:
    PSUTIL_AVAILABLE = False
    logger.warning("psutil is not available. Fallback method will be used for terminating Word processes.")

def create_word_letter(letter_content: str, output_path: Path, company_footer: str = COMPANY_FOOTER) -> None:
    """Create a Word (DOCX) letter with styling and formatting."""
    try:
        doc = Document()

        # Configure page
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.top_margin = Cm(1.54)
        section.bottom_margin = Cm(0)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

        # Set style
        style = doc.styles['Normal']
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

        # Add logo
        logo_path = Path(get_asset_path("asset/derland2.png"))
        if logo_path.exists():
            header = section.header
            header.is_linked_to_previous = False
            header_paragraph = header.paragraphs[0]
            header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            header_run = header_paragraph.add_run()
            header_run.add_picture(str(logo_path), width=Inches(1.93), height=Inches(0.55))

        # Add content
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("")

        lines = letter_content.splitlines()
        sign_flg = False
        sign_cnt = 0
        for line in lines:
            if "Re:" in line:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(line)
                run.bold = True
                continue

            if "The amount being" in line or "Please note that" in line:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(line)
                run.bold = True
                continue

            if any(x in line for x in ["1)", "2)", "3)"]):
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(line)
                run.italic = True
                continue

            if "which ALL" in line:
                paragraph = doc.add_paragraph()
                before, after = line.split("ALL", 1)
                paragraph.add_run(before)
                run = paragraph.add_run("ALL")
                run.bold = True
                run.underline = True
                paragraph.add_run(after)
                continue

            if sign_flg:
                sign_cnt += 1

            if "Yours sincerely," in line:
                sign_flg = True

            if sign_cnt == 2:
                signature_path = Path(get_asset_path("asset/sign.png"))
                if signature_path.exists():
                    sig_paragraph = doc.add_paragraph()
                    sig_run = sig_paragraph.add_run()
                    sig_run.add_picture(str(signature_path), width=Inches(0.9), height=Inches(0.63))

            if sign_cnt < 3 or sign_cnt >= 6:
                if "DARLANDS" in line:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(line)
                    run.bold = True
                else:
                    doc.add_paragraph(line)

        # Add footer
        section.footer_distance = Cm(0)
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.text = ""
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for line in company_footer.split('\n'):
            run = footer_paragraph.add_run(line + "\n")
            run.font.name = "Calibri"
            run.font.size = Pt(11)

        doc.save(str(output_path))

    except Exception as e:
        logger.error(f"Error creating Word DOCX: {e}")
        raise GenerationError(f"Error creating Word DOCX: {str(e)}")

def terminate_word_processes():
    """Forcefully terminate all running Microsoft Word processes."""
    if PSUTIL_AVAILABLE:
        for proc in psutil.process_iter(['name']):
            if proc.info['name'] == 'WINWORD.EXE':
                try:
                    proc.terminate()
                    logger.info(f"Terminated Word process with PID {proc.pid}")
                except Exception as e:
                    logger.error(f"Failed to terminate Word process with PID {proc.pid}: {e}")
    else:
        # Fallback method using subprocess
        try:
            subprocess.run(["taskkill", "/F", "/IM", "WINWORD.EXE"], check=True, capture_output=True)
            logger.info("Terminated Word processes using fallback method")
        except subprocess.CalledProcessError as e:
            logger.error(f"Failed to terminate Word processes using fallback method: {e}")

def convert_pdf_letter(letter_content: str, output_path: Path) -> None:
    """Convert Word letter to PDF with retry logic and fallback methods."""
    try:
        # Create Word document with same name as final PDF
        docx_path = output_path.with_suffix('.docx')
        
        # Create Word document
        create_word_letter(letter_content, docx_path)
        
        logger.info(f"Word document created successfully: {docx_path}")
        
        # Check if the output directory exists and is writable
        output_dir = output_path.parent
        if not output_dir.exists():
            logger.error(f"Output directory does not exist: {output_dir}")
            raise GenerationError(f"Output directory does not exist: {output_dir}")
        if not os.access(str(output_dir), os.W_OK):
            logger.error(f"No write permission for output directory: {output_dir}")
            raise GenerationError(f"No write permission for output directory: {output_dir}")
        
        # Attempt PDF conversion with retry logic
        max_retries = 3
        for attempt in range(max_retries):
            try:
                logger.info(f"Attempting to convert {docx_path} to {output_path} (Attempt {attempt + 1})")
                convert(str(docx_path), str(output_path))
                logger.info(f"PDF conversion successful: {output_path}")
                return
            except Exception as e:
                logger.warning(f"PDF conversion attempt {attempt + 1} failed: {str(e)}")
                terminate_word_processes()
                time.sleep(2)  # Wait before retrying
        
        # If all attempts fail, try fallback method
        logger.error("All conversion attempts failed. Trying fallback method.")
        # fallback_convert_pdf(docx_path, output_path)
        
        # Check if the PDF was actually created
        if not output_path.exists():
            logger.error(f"PDF file was not created: {output_path}")
            raise GenerationError(f"PDF file was not created: {output_path}")
        
    except Exception as e:
        logger.error(f"Error creating PDF: {str(e)}", exc_info=True)
        raise GenerationError(f"Error creating PDF: {str(e)}")

def fallback_convert_pdf(docx_path: Path, output_path: Path) -> None:
    """Fallback method to convert DOCX to PDF using ReportLab."""
    try:
        # Read the content from the DOCX file
        doc = Document(docx_path)
        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        # Create a PDF using ReportLab
        pdf = SimpleDocTemplate(str(output_path), pagesize=letter)
        styles = getSampleStyleSheet()
        flowables = []

        for line in content.split('\n'):
            p = Paragraph(line, styles['Normal'])
            flowables.append(p)
            flowables.append(Spacer(1, 12))

        pdf.build(flowables)
        logger.info(f"Fallback PDF conversion successful: {output_path}")
    except Exception as e:
        logger.error(f"Fallback PDF conversion failed: {str(e)}")
        raise GenerationError(f"Fallback PDF conversion failed: {str(e)}")

def generate_letter(
    content: str,
    letter_type: str = "annual",
    page_count: int = 1,
    override_names: Optional[str] = None,
    override_address: Optional[Dict] = None,
    override_salutation_name: Optional[str] = None
) -> tuple:
    """
    Generate a wayleave letter based on document content.
    
    Args:
        content: The document content to process
        letter_type: Type of letter ('annual' or '15-year')
        page_count: Number of pages in the document
        override_names: Optional names to use instead of extracted ones
        override_address: Optional address dict to use instead of extracted one
        override_salutation_name: Optional salutation name to use in the Dear {} section
    """
    try:
        logger.info(f"Generating {letter_type} letter")
        
        if letter_type not in ["annual", "15-year"]:
            raise GenerationError(f"Invalid letter type: {letter_type}")
        
        # Extract information
        if letter_type == "annual":
            info = extract_names_and_address_annual(content)
            template = ANNUAL_LETTER_TEMPLATE
        else:
            info = extract_names_and_address_fifteen_year(content)
            template = FIFTEEN_YEAR_LETTER_TEMPLATE
        
        # Use override values if provided
        names_to_use = override_names if override_names is not None else info['full_names']
        address_to_use = override_address if override_address is not None else info['address']
        
        # Format details
        header_names, _ = format_names(names_to_use)  # Get header names only
        formatted_address = format_address(address_to_use)
        current_date = datetime.now().strftime("%d %B %Y")
        page_counts = (page_count - 1) if letter_type == "annual" else page_count
        sign_page = num2words(page_counts, to="ordinal").upper()

        # Use override_salutation_name directly if provided, otherwise use formatted salutation
        salutation_names = override_salutation_name if override_salutation_name is not None else info['salutation_name']

        # Generate letter
        letter = template.format(
            current_date,
            f"{header_names}\n{formatted_address}",
            salutation_names,  # Use salutation name directly
            sign_page
        )
        
        filename = generate_filename(address_to_use)
        
        logger.info("Letter generation successful")
        return letter, filename
        
    except Exception as e:
        logger.error(f"Error generating {letter_type} letter: {e}")
        raise GenerationError(f"Error generating letter: {str(e)}")

def generate_second_letter(
    content: str,
    letter_type: str = "annual",
    override_names: Optional[str] = None,
    override_address: Optional[Dict] = None,
    override_salutation_name: Optional[str] = None
) -> tuple:
    """
    Generate a second wayleave letter based on document content.
    
    Args:
        content: The document content to process
        letter_type: Type of letter ('annual' or '15-year')
        override_names: Optional names to use instead of extracted ones
        override_address: Optional address dict to use instead of extracted one
        override_salutation_name: Optional salutation name to use in the Dear {} section
    """
    try:
        logger.info(f"Generating second {letter_type} letter")
        
        if letter_type not in ["annual", "15-year"]:
            raise GenerationError(f"Invalid letter type: {letter_type}")
        
        # Extract information
        if letter_type == "annual":
            info = extract_names_and_address_annual(content)
        else:
            info = extract_names_and_address_fifteen_year(content)
        
        # Use override values if provided
        names_to_use = override_names if override_names is not None else info['full_names']
        address_to_use = override_address if override_address is not None else info['address']
        
        # Format details
        header_names, _ = format_names(names_to_use)  # Get header names only
        formatted_address = format_address(address_to_use)
        current_date = datetime.now().strftime("%d %B %Y")

        # Use override_salutation_name directly if provided, otherwise use formatted salutation
        salutation_names = override_salutation_name if override_salutation_name is not None else info['salutation_name']

        # Generate letter
        letter = SECOND_LETTER_TEMPLATE.format(
            current_date,
            f"{header_names}\n{formatted_address}",
            salutation_names,  # Use salutation name directly
        )
        
        filename = generate_filename(address_to_use)
        
        logger.info("Second letter generation successful")
        return letter, filename
        
    except Exception as e:
        logger.error(f"Error generating second {letter_type} letter: {e}")
        raise GenerationError(f"Error generating second letter: {str(e)}")