"""Module for generating wayleave letters."""
from datetime import datetime
import re
import os
from pathlib import Path
import logging
import fitz  # PyMuPDF
from num2words import num2words
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENTATION
from docx2pdf import convert
from constants import get_asset_path
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch

# Configure logging
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class ContentError(Exception):
    """Exception raised for content-related errors."""
    pass

class WayleaveLetterGenerator:
    """Class for generating wayleave letters."""
    
    def __init__(self):
        """Initialize the letter generator with templates."""
        self.company_footer = """Autaway Ltd t.a Darlands   Suite 2063 6-8 Revenge Road Lordswood Kent ME5 8UD
E: info@darlands.co.uk     W: darlands.co.uk     Company No. 12185075"""
        
        self.annual_letter_template = """{}

{}


Dear {}

Re: Electrical Equipment on your Land – Wayleave Agreement

We are pleased to inform you that we have now secured agreement for payment to be made to you from Scottish & Southern Energy (SSE).

Please find enclosed two copies of your Wayleave agreement which ALL registered homeowners must sign. These documents confirm that SSE hold electrical equipment on your land and as such they will now make a wayleave payment to you.

The amount being offered to you is confirmed on the agreement under 'Section 1: the Wayleave Payment'.

To help you complete the agreement, please follow these steps for both copies of the wayleave agreements:

        1) All homeowners must sign on the {} PAGE
        2) All homeowners must sign and date on the FOURTH PAGE (Title Plan)
        3) Send both copies back to us in the prepaid envelope provided

Please note that there is no cost, or charge to you whatsoever for us setting your wayleave up. All the monies for the wayleave will be paid to, and kept by you.

Yours sincerely,






Paul Wakeford
Partner
DARLANDS"""

        self.fifteen_year_letter_template = """{}

{}


Dear {}

Re: Electrical Equipment on your Land – Wayleave Agreement

We are pleased to inform you that we have now secured agreement for payment to be made to you from Scottish & Southern Energy (SSE).

Please find enclosed two copies of your Wayleave agreement which ALL registered homeowners must sign. These documents confirm that SSE hold electrical equipment on your land and as such they will now make a wayleave payment to you.

The amount being offered to you is confirmed on the agreement under 'Section 1: the Wayleave Payment'.

To help you complete the agreement, please follow these steps for both copies of the wayleave agreements:

        1) All homeowners must sign on the {} PAGE
        2) All homeowners must sign and date on the FOURTH PAGE (Title Plan)
        3) Send both copies back to us in the prepaid envelope provided

Please note that there is no cost, or charge to you whatsoever for us setting your wayleave up. All the monies for the wayleave will be paid to, and kept by you. 

Yours sincerely,






Paul Wakeford
Partner
DARLANDS"""

        self.second_letter_template = """{}

{}


Dear {}

Re: Completed Wayleave Enclosed

I am pleased to enclose your countersigned wayleave with SSE accompanied by the cheque payment from them. This now completes the wayleave process for you and we will therefore close our files.

It has been a pleasure representing you in this matter.

Yours sincerely,






Paul Wakeford	
Partner
DARLANDS
"""

    def create_word_letter(self, letter_content: str, output_path: Path) -> None:
        """
        Create a Word (DOCX) letter with styling and formatting, set to A4 size,
        with reduced vertical spacing to minimize tall highlights or bounding areas.
        """
        try:
            doc = Document()

            # 1) Configure page size, orientation, and margins (adjust as desired)
            section = doc.sections[0]
            section.page_width = Cm(21.0)    # A4
            section.page_height = Cm(29.7)   # A4
            section.orientation = WD_ORIENTATION.PORTRAIT
            section.top_margin = Cm(1.54)
            section.bottom_margin = Cm(0)
            section.left_margin = Cm(2.54)
            section.right_margin = Cm(2.54)

            # 2) Set the default (Normal) style, with tighter line spacing
            style = doc.styles['Normal']
            style.font.name = "Calibri"
            style.font.size = Pt(11)

            paragraph_format = style.paragraph_format
            # Force single line spacing
            paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            # Optionally reduce or remove extra space before/after each paragraph:
            paragraph_format.space_before = Pt(0)
            paragraph_format.space_after = Pt(0)  # or Pt(0) if you want even less

            # 3) Header with optional logo
            logo_path = Path(get_asset_path("asset/derland.png"))
            if logo_path.exists():
                header = section.header
                header.is_linked_to_previous = False
                header_paragraph = header.paragraphs[0]
                header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                header_run = header_paragraph.add_run()
                # Resize the logo as needed
                header_run.add_picture(str(logo_path), width=Inches(1.93), height=Inches(0.55))

            # 4) Letter content, line by line
            doc.add_paragraph("")
            doc.add_paragraph("")
            doc.add_paragraph("")
            doc.add_paragraph("")
            lines = letter_content.splitlines()
            sign_flg = False
            sign_cnt = 0
            for line in lines:
                # Each line becomes its own paragraph, including blank ones

                if "Re:" in line:
                    # Example: Make the entire line bold
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(line)
                    run.bold = True
                    continue

                if "The amount being" in line:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(line)
                    run.bold = True
                    continue

                if "Please note that" in line:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(line)
                    run.bold = True
                    continue
                
                if "1)" in line:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(line)
                    run.italic = True
                    continue

                if "2)" in line:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(line)
                    run.italic = True
                    continue

                if "3)" in line:
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run(line)
                    run.italic = True
                    continue

                if "which ALL" in line:
                    paragraph = doc.add_paragraph()
                    before, after = line.split("ALL", 1)
                    paragraph.add_run(before)
                    run = paragraph.add_run("ALL")
                    run.bold = True
                    run.underline = True
                    paragraph.add_run(after)
                    continue

                if sign_flg == True:
                    sign_cnt = sign_cnt + 1

                if "Yours sincerely," in line:
                    sign_flg = True

                if sign_cnt == 2:
                    # 5) Optional signature
                    signature_path = Path(get_asset_path("asset/sign.png"))
                    if signature_path.exists():
                        sig_paragraph = doc.add_paragraph()
                        sig_run = sig_paragraph.add_run()
                        sig_run.add_picture(str(signature_path), width=Inches(0.9), height=Inches(0.63))

                if sign_cnt < 3 or sign_cnt >= 6:
                    if "DARLANDS" in line:
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run(line)
                        run.bold = True
                    else:
                        doc.add_paragraph(line)

            # 6) Footer text
            section.footer_distance = Cm(0)  # Increase as needed
            footer = section.footer
            footer.is_linked_to_previous = False
            footer_paragraph = footer.paragraphs[0]
            footer_paragraph.text = ""
            footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for line in self.company_footer.split('\n'):
                run = footer_paragraph.add_run(line + "\n")
                run.font.name = "Calibri"
                run.font.size = Pt(11)

            # 7) Save the Word document
            doc.save(str(output_path))

        except Exception as e:
            logger.error(f"Error creating Word DOCX: {e}")
            raise ContentError(f"Error creating Word DOCX: {str(e)}")

    def convert_pdf_letter(self, letter_content: str, output_path: Path) -> None:
        try:
            # Create a temporary DOCX file
            docx_path = output_path.with_suffix('.docx')
            
            # Create the Word document
            self.create_word_letter(letter_content, docx_path)
            
            # Convert DOCX to PDF
            convert(str(docx_path), str(output_path))
            
            # Clean up the temporary DOCX file
            if docx_path.exists():
                docx_path.unlink()
                
        except Exception as e:
            # Clean up temporary file if it exists
            if 'docx_path' in locals() and docx_path.exists():
                docx_path.unlink()
            raise ContentError(f"Error creating PDF: {str(e)}")

    def create_pdf_letter(self, letter_content: str, output_path: Path) -> None:
        """
        Create a PDF letter with proper formatting.
        
        Args:
            letter_content: The content of the letter
            output_path: Path where to save the PDF
        """
        try:
            # Use get_asset_path to get the correct absolute paths for images
            logo_path = get_asset_path("asset/derland.png")
            signature_path = get_asset_path("asset/sign.png")

            # Create a new PDF document
            doc = SimpleDocTemplate(str(output_path), pagesize=A4, rightMargin=72, leftMargin=72, topMargin=40, bottomMargin=20)
            
            # Set margins (in points, 1 inch = 72 points)
            left_margin = 72
            top_margin = 72
            right_margin = 72
            
            styles = getSampleStyleSheet()

            # 3. Customize a Paragraph Style
            #    - Use 'Normal' as a base and override line spacing and font size, etc.
            style_normal = styles["Normal"]
            style_normal.fontName = "Helvetica"
            style_normal.fontSize = 11
            style_normal.leading = 16  # line spacing in points
            style_normal.spaceAfter = 0  # extra space after each paragraph

            story = []

            logo_img = Image(logo_path, width=120, height=32, hAlign="LEFT")
            signature_img = Image(signature_path, width=60, height=45, hAlign="LEFT")
            story.append(logo_img)

            story.append(Spacer(1, 10))
            story.append(Spacer(1, 10))
            story.append(Spacer(1, 10))
            
            paragraphs = letter_content.split("\n")

            sign_flg = False
            sign_cnt = 0
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para, style_normal))
                elif sign_flg == True and (sign_cnt < 2 or sign_cnt > 3):
                    logger.info(f"skiping line {sign_cnt}")
                else:
                    story.append(Spacer(1, 10))

                if sign_flg == True:
                    sign_cnt = sign_cnt + 1

                if "Yours sincerely" in para:
                    sign_flg = True

                if sign_cnt == 3:
                    story.append(signature_img)
            
            def draw_footer(canvas, doc_):
                """
                Draws a multi-line footer using a Paragraph, converting newlines to <br/>.
                """
                canvas.saveState()

                # Convert Python newlines to <br/> so Paragraph displays them as multiple lines
                footer_html = self.company_footer.replace("\n", "<br/>")

                footer_style = ParagraphStyle(
                    name="FooterStyle",
                    parent=styles["Normal"],
                    fontName="Helvetica",
                    fontSize=10,
                    leading=11,
                    alignment=1,  # center
                    spaceAfter=10
                )

                footer_para = Paragraph(footer_html, footer_style)

                # Wrap the paragraph to find its height
                max_width = doc_.width
                w, h = footer_para.wrapOn(canvas, max_width, doc_.bottomMargin)

                # Decide where to draw it (for example, 0.5 inch from bottom)
                x_pos = doc_.leftMargin
                y_pos = 0.5 * inch

                # Draw it
                footer_para.drawOn(canvas, x_pos, y_pos)

                canvas.restoreState()

            # -- 5) Build the PDF with our on-page functions
            #       onFirstPage -> runs on page 1
            #       onLaterPages -> runs on pages 2,3,4...
            doc.build(
                story, 
                onFirstPage=draw_footer,
                onLaterPages=draw_footer
            )
        except Exception as e:
            logger.error(f"Error creating PDF: {e}")
            raise ContentError(f"Error creating PDF: {str(e)}")

    def extract_names_and_address_annual(self, content: str) -> dict:
        """Extract names and address from annual wayleave document content (updated to handle multiple postcodes)."""
        try:
            logger.debug("Raw content:")
            logger.debug("-" * 80)
            logger.debug(content)
            logger.debug("-" * 80)

            self.validate_content(content, "annual")

            logger.debug("Extracting names from annual document")
            name_match = re.search(r'I/We,\s+([^\n]+)', content)
            if not name_match:
                logger.error("Name pattern not found in content")
                logger.debug(f"Content preview: {content[:200]}")
                raise ContentError("Could not find names in document")

            names = name_match.group(1).strip()
            logger.debug(f"Found names: {names}")

            logger.debug("Extracting address from annual document")
            address_match = re.search(r'of\s+(.*?)(?=being|$)', content, re.DOTALL)
            if not address_match:
                logger.error("Address pattern not found in content")
                raise ContentError("Could not find address starting with 'of'")

            address = address_match.group(1).strip()
            logger.debug(f"Found raw address: {address}")

            # --------------------------------------------------
            # Updated: find ALL possible postcodes in the address
            # --------------------------------------------------
            postcode_pattern = r'[A-Z]{1,2}\d[\dA-Z]?\s*\d[A-Z]{2}'
            multiple_postcodes = re.findall(postcode_pattern, address, flags=re.IGNORECASE)
            if not multiple_postcodes:
                logger.error("Postcode not found in address")
                raise ContentError("Could not find valid postcode in address")

            logger.debug(f"Found multiple postcodes: {multiple_postcodes}")

            # Use the first matched postcode as the "primary" one
            primary_postcode = multiple_postcodes[0].upper()

            # Locate this first postcode in the address string so we can separate it out
            first_pc_index = address.upper().find(primary_postcode)
            if first_pc_index == -1:
                logger.error("Could not find the location of the first postcode in the address")
                raise ContentError("Could not parse the address properly")

            # Slice the address up to the start of that postcode
            address_without_postcode = address[:first_pc_index].strip()
            if address_without_postcode.endswith(","):
                address_without_postcode = address_without_postcode[:-1]

            parts = [p.strip() for p in address_without_postcode.split(",")]
            logger.debug(f"Address parts: {parts}")

            if len(parts) < 2:
                logger.error("Not enough address parts found")
                raise ContentError("Address format is incomplete")

            # Build the final result
            result = {
                'full_names': names,
                'address': {
                    'house': parts[0],
                    'city': parts[1],
                    'county': parts[2] if len(parts) > 2 else '',
                    'postcode': " ".join([pc.upper() for pc in multiple_postcodes]),
                    # Store all postcodes in a list if desired:
                    'all_postcodes': [pc.upper() for pc in multiple_postcodes]
                }
            }

            logger.debug(f"Successfully extracted information: {result}")
            return result

        except ContentError:
            raise
        except Exception as e:
            logger.error(f"Error extracting names and address from annual document: {e}")
            raise ContentError(f"Error processing annual document: {str(e)}")

    def extract_names_and_address_fifteen_year(self, content: str) -> dict:
        """Extract names and address from 15-year wayleave document content."""
        try:
            logger.debug("Raw content:")
            logger.debug("-" * 80)
            logger.debug(content)
            logger.debug("-" * 80)
            
            # 1) Extract the "names" and "house_and_road"
            #    (same as before)
            pattern = r'\(1\)\s*(.*?)\s+of\s*(.*?)\s*(?=,\s*|\))'
            match = re.search(pattern, content, re.DOTALL)
            
            if not match:
                logger.error("Name and address pattern not found in content")
                logger.debug(f"Content preview: {content[:500]}")
                raise ContentError("Could not find names and address in 15-year document")
            
            names = match.group(1).strip()
            house_and_road = match.group(2).strip()
            house_and_road = house_and_road.replace('\n', '')
            # 2) Updated pattern to allow 2 or 3 commas before postcode
            address_pattern = (
                r',\s*([^,]+),'         # group(1)
                r'\s*([^,]+)'          # group(2)
                r'(?:,\s*([^,]+))?'    # optional group(3)
                r'\s+([A-Z0-9][A-Z0-9\s-]{0,10}[A-Z0-9])'  # group(4) = postcode
            )
            address_match = re.search(address_pattern, content)
            
            if not address_match:
                # Optional fallback or error:
                logger.error("Could not find address details with 2 or 3 commas + postcode")
                raise ContentError("Could not find complete address details")
            
            # 3) Extract the captures
            line1 = address_match.group(1).strip()  # e.g., 'Highworth' or 'Lightwater'
            line2 = address_match.group(2).strip()  # e.g., 'Swindon' or 'Surrey'
            line3 = address_match.group(3).strip() if address_match.group(3) else ''
            postcode_match = address_match.group(4).strip()  # single postcode from group(4)
            
            if '(' in line2:
                parts = line2.split()
                if parts:
                    line2 = parts[0]

            # 4) The entire substring matched by this pattern
            address_text = address_match.group(0)
            
            # Also find *all* postcodes in that substring, just in case there’s more than one
            postcode_pattern = r'[A-Z]{1,2}\d[\dA-Z]?\s*\d[A-Z]{2}'
            multiple_postcodes = re.findall(postcode_pattern, address_text, flags=re.IGNORECASE)
            
            # If multiple postcodes are found, pick the first for backward compatibility
            postcode = multiple_postcodes[0].upper() if multiple_postcodes else postcode_match
            
            logger.debug(f"Found names: {names}")
            logger.debug(f"house_and_road: {house_and_road}")
            logger.debug(f"line1: {line1}")
            logger.debug(f"line2: {line2}")
            logger.debug(f"line3: {line3}")
            logger.debug(f"postcode: {postcode}")
            
            # 5) Clean up the name string
            names = re.sub(r'\s+', ' ', names)
            
            # 6) Split house_and_road for backward compatibility
            house_parts = house_and_road.split(' ', 1)
            house = house_parts[0]
            road = house_parts[1] if len(house_parts) > 1 else ''
            
            logger.info(f"multiple_postcodes => {multiple_postcodes}")
            
            logger.info(f"house and road: {house_and_road}, city:: {line1},  county:: {line2}, and multipostcodes: {" ".join([pc.upper() for pc in multiple_postcodes])}")
            # Build final result object
            result = {
                'full_names': names,
                'address': {
                    'house': house_and_road,   # or house_and_road
                    'road': road,
                    'house_and_road': house_and_road,
                    # line1, line2, line3 might be city / county / extra
                    'city': line1,
                    'county': line2,
                    'line1': line1,
                    'line2': line2,
                    'line3': line3,
                    'postcode': " ".join([pc.upper() for pc in multiple_postcodes]),
                    # Optionally store all matched postcodes
                    'all_postcodes': [pc.upper() for pc in multiple_postcodes],
                }
            }
            
            logger.debug(f"Successfully extracted information: {result}")
            return result
            
        except ContentError:
            raise
        except Exception as e:
            logger.error(f"Error extracting names and address from 15-year document: {e}")
            raise ContentError(f"Error processing 15-year document: {str(e)}")

    def validate_content(self, content: str, letter_type: str) -> None:
        """Validate the content before processing."""
        if not content or not isinstance(content, str):
            raise ContentError("Document content is empty or invalid")
            
        content = content.strip()
        if not content:
            raise ContentError("Document content is empty after trimming whitespace")
            
        if len(content) < 50:
            raise ContentError("Document content is too short to be valid")
            
        if letter_type == "annual":
            if "ELECTRICITY ACT 1989" not in content and "Re: Electrical Equipment" not in content:
                raise ContentError("Content does not appear to be a valid annual wayleave document")
        elif letter_type == "15-year":
            if "This Agreement" not in content and "AGREED TERMS" not in content:
                raise ContentError("Content does not appear to be a valid 15-year wayleave document")
                
        logger.debug(f"Content validation passed for {letter_type} document")

    def format_names(self, full_names: str) -> tuple:
        """Format names for header and salutation."""
        try:
            if not full_names:
                raise ContentError("Names string is empty")
                
            header_names = full_names.replace(' AND ', ' & ').replace(' and ', ' & ')
            
            first_names = []
            for name in re.split(' AND | & ', full_names):
                name = name.strip()
                if not name:
                    continue
                name_parts = name.split()
                if not name_parts:
                    continue
                first_names.append(name_parts[0])
            
            if not first_names:
                raise ContentError("No valid names found")
            
            if len(first_names) == 2:
                salutation_names = f"{first_names[0]} and {first_names[1]}"
            elif len(first_names) > 2:
                salutation_names = ", ".join(first_names[:-1]) + f", and {first_names[-1]}"
            else:
                salutation_names = first_names[0]
                
            return header_names, salutation_names
            
        except Exception as e:
            logger.error(f"Error formatting names: {e}")
            raise ContentError(f"Error formatting names: {str(e)}")

    def format_address(self, address_dict: dict) -> str:
        """Format address dictionary into string with proper line breaks."""
        try:
            if not isinstance(address_dict, dict):
                raise ContentError("Invalid address dictionary")
                
            address_parts = [
                address_dict['house'],
                address_dict['city'],
                'Surrey' if 'county' in address_dict and 'Surrey' in address_dict['county'] else address_dict.get('county', 'Surrey'),
                address_dict['postcode']
            ]
            
            return '\n'.join(part for part in address_parts if part)
            
        except Exception as e:
            logger.error(f"Error formatting address: {e}")
            raise ContentError(f"Error formatting address: {str(e)}")

    def generate_filename(self, address_dict: dict) -> str:
        """Generate filename from address components."""
        try:
            if not isinstance(address_dict, dict):
                raise ContentError("Invalid address dictionary")
            
            # Clean and sanitize each component
            def sanitize_component(component):
                if not component:
                    return ""
                # Replace newlines and multiple spaces with a single space
                component = re.sub(r'\s+', ' ', component.strip())
                # Replace invalid filename characters
                component = re.sub(r'[<>:"/\\|?*]', '', component)
                return component
            
            filename_parts = [
                sanitize_component(address_dict['house']),
                sanitize_component(address_dict['city']),
                sanitize_component('Surrey' if 'county' in address_dict and 'Surrey' in address_dict['county'] else address_dict.get('county', 'Surrey')),
                sanitize_component(address_dict['postcode'])
            ]
            
            # Filter out empty parts and join with commas
            filename = ", ".join(part for part in filename_parts if part)
            
            # Ensure the filename ends with .pdf
            if not filename.lower().endswith('.pdf'):
                filename += '.pdf'
                
            return filename
            
        except Exception as e:
            logger.error(f"Error generating filename: {e}")
            raise ContentError(f"Error generating filename: {str(e)}")

    def generate_letter(self, content: str, letter_type: str = "annual", page_count: int = 1) -> tuple:
        """Generate a wayleave letter based on document content."""
        try:
            logger.info(f"Generating {letter_type} letter")
            logger.debug(f"Content preview: {content[:200] if content else 'Empty content'}")
            
            if letter_type not in ["annual", "15-year"]:
                raise ContentError(f"Invalid letter type: {letter_type}")
            
            if letter_type == "annual":
                info = self.extract_names_and_address_annual(content)
                template = self.annual_letter_template
            else:
                info = self.extract_names_and_address_fifteen_year(content)
                template = self.fifteen_year_letter_template
            
            header_names, salutation_names = self.format_names(info['full_names'])
            formatted_address = self.format_address(info['address'])
            
            current_date = datetime.now().strftime("%d %B %Y")
            
            page_counts = (page_count - 1) if letter_type == "annual" else page_count

            sign_page = num2words(page_counts, to="ordinal").upper()

            letter = template.format(
                current_date,
                f"{header_names}\n{formatted_address}",
                salutation_names,
                sign_page
            )
            
            filename = self.generate_filename(info['address'])
            
            logger.info("Letter generation successful")
            return letter, filename
            
        except Exception as e:
            logger.error(f"Error generating {letter_type} letter: {e}")
            raise ContentError(f"Error generating letter: {str(e)}")

    def generate_second_letter(self, content: str, letter_type: str = "annual", page_count: int = 1) -> tuple:
        """Generate a wayleave letter based on document content."""
        try:
            logger.info(f"Generating {letter_type} letter")
            logger.debug(f"Content preview: {content[:200] if content else 'Empty content'}")
            
            if letter_type not in ["annual", "15-year"]:
                raise ContentError(f"Invalid letter type: {letter_type}")
            
            if letter_type == "annual":
                info = self.extract_names_and_address_annual(content)
            else:
                info = self.extract_names_and_address_fifteen_year(content)
            
            template = self.second_letter_template
            header_names, salutation_names = self.format_names(info['full_names'])
            formatted_address = self.format_address(info['address'])
            
            current_date = datetime.now().strftime("%d %B %Y")
            
            letter = template.format(
                current_date,
                f"{header_names}\n{formatted_address}",
                salutation_names,
            )
            
            filename = self.generate_filename(info['address'])
            
            logger.info("Letter generation successful")
            return letter, filename
            
        except Exception as e:
            logger.error(f"Error generating {letter_type} letter: {e}")
            raise ContentError(f"Error generating letter: {str(e)}")